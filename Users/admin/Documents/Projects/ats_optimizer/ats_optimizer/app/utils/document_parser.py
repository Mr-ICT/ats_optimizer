"""
app/utils/document_parser.py
Robust document parser for PDF and DOCX files.
Returns extracted text, layout metadata, and a parsing confidence score.
"""
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """Container for parsed document data."""
    text: str = ""
    pages: int = 0
    confidence: float = 0.0          # 0.0–1.0
    warnings: list = field(default_factory=list)
    layout_flags: dict = field(default_factory=dict)
    word_count: int = 0
    unreadable_sections: list = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "text": self.text,
            "pages": self.pages,
            "confidence": round(self.confidence, 2),
            "confidence_pct": round(self.confidence * 100, 1),
            "warnings": self.warnings,
            "layout_flags": self.layout_flags,
            "word_count": self.word_count,
            "unreadable_sections": self.unreadable_sections,
        }


class DocumentParser:
    """
    Parses PDF and DOCX files into clean text with metadata.
    Strategy:
      - PDF  → PyMuPDF (fitz); fallback warns if page is image-only
      - DOCX → python-docx; extracts paragraphs + table cells
    """

    def parse(self, file_path: str) -> ParseResult:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        else:
            result = ParseResult()
            result.warnings.append(f"Unsupported file type: {ext}")
            return result

    # ─── PDF Parsing ────────────────────────────────────────────────────────

    def _parse_pdf(self, path: Path) -> ParseResult:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError("PyMuPDF not installed. Run: pip install PyMuPDF")

        result = ParseResult()
        full_text_parts = []
        total_chars = 0
        readable_chars = 0

        try:
            doc = fitz.open(str(path))
            result.pages = len(doc)

            for page_num, page in enumerate(doc):
                # Extract text blocks preserving layout
                blocks = page.get_text("blocks")  # list of (x0,y0,x1,y1,text,bno,btype)

                page_text_parts = []
                page_char_count = 0

                for block in blocks:
                    if len(block) >= 5:
                        block_text = block[4].strip()
                        block_type = block[6] if len(block) > 6 else 0

                        if block_type == 0 and block_text:  # 0 = text block
                            page_text_parts.append(block_text)
                            page_char_count += len(block_text)

                page_text = "\n".join(page_text_parts)

                # Detect image-heavy / unreadable pages
                if page_char_count < 50:
                    result.unreadable_sections.append(f"Page {page_num + 1} appears image-only or unreadable")
                    result.warnings.append(f"Page {page_num + 1}: low text content ({page_char_count} chars)")
                else:
                    readable_chars += page_char_count

                total_chars += page_char_count
                full_text_parts.append(page_text)

            doc.close()

            result.text = self._clean_text("\n\n".join(full_text_parts))
            result.word_count = len(result.text.split())

            # Layout flags — detect common ATS-unfriendly patterns
            result.layout_flags = self._detect_layout_flags_pdf(path)

            # Confidence: ratio of readable chars to total, capped at 1.0
            if total_chars > 0:
                result.confidence = min(readable_chars / total_chars, 1.0)
                if result.word_count < 50:
                    result.confidence *= 0.5
                    result.warnings.append("Very little text extracted — resume may be image-based")
            else:
                result.confidence = 0.0
                result.warnings.append("No text could be extracted from PDF")

        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            result.warnings.append(f"PDF parsing failed: {str(e)}")
            result.confidence = 0.0

        return result

    def _detect_layout_flags_pdf(self, path: Path) -> dict:
        """Detect multi-column, tables, text boxes in PDF."""
        flags = {
            "multi_column": False,
            "has_tables": False,
            "has_graphics": False,
            "has_headers_footers": False,
        }
        try:
            import fitz
            doc = fitz.open(str(path))
            for page in doc:
                blocks = page.get_text("blocks")
                # Heuristic: if blocks span very different x-positions → multi-column
                x_positions = [b[0] for b in blocks if len(b) >= 5 and b[6] == 0]
                if x_positions:
                    x_range = max(x_positions) - min(x_positions)
                    page_width = page.rect.width
                    if x_range > page_width * 0.35:
                        flags["multi_column"] = True

                # Check for image/drawing objects
                image_list = page.get_images()
                if image_list:
                    flags["has_graphics"] = True

            doc.close()
        except Exception:
            pass
        return flags

    # ─── DOCX Parsing ───────────────────────────────────────────────────────

    def _parse_docx(self, path: Path) -> ParseResult:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        result = ParseResult()
        text_parts = []

        try:
            doc = Document(str(path))

            # Extract all paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

            # Extract table cell text (ATS often can't read tables)
            table_texts = []
            for table in doc.tables:
                result.layout_flags["has_tables"] = True
                result.warnings.append("Document contains tables — ATS may not read table content correctly")
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            table_texts.append(cell_text)

            # Combine paragraph + table text
            all_text = "\n".join(text_parts)
            if table_texts:
                all_text += "\n" + "\n".join(table_texts)

            result.text = self._clean_text(all_text)
            result.word_count = len(result.text.split())
            result.pages = max(1, result.word_count // 300)  # estimate

            # Layout detection
            if not result.layout_flags.get("has_tables"):
                result.layout_flags["has_tables"] = False
            result.layout_flags["multi_column"] = self._detect_docx_columns(doc)
            result.layout_flags["has_graphics"] = len(doc.inline_shapes) > 0
            result.layout_flags["has_headers_footers"] = self._has_headers_footers(doc)

            if result.layout_flags["has_graphics"]:
                result.warnings.append("Document contains inline images/graphics — avoid for ATS")
            if result.layout_flags["multi_column"]:
                result.warnings.append("Multi-column layout detected — ATS parsers may misread column order")
            if result.layout_flags["has_headers_footers"]:
                result.warnings.append("Headers/footers detected — ATS may ignore or duplicate this content")

            # Confidence based on word count
            if result.word_count >= 150:
                result.confidence = 0.95
            elif result.word_count >= 50:
                result.confidence = 0.75
            else:
                result.confidence = 0.40
                result.warnings.append("Very little text extracted from DOCX")

        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            result.warnings.append(f"DOCX parsing failed: {str(e)}")
            result.confidence = 0.0

        return result

    def _detect_docx_columns(self, doc) -> bool:
        """Heuristic: check section column settings."""
        try:
            for section in doc.sections:
                cols = section._sectPr.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
                )
                for col in cols:
                    num = col.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
                    )
                    if num and int(num) > 1:
                        return True
        except Exception:
            pass
        return False

    def _has_headers_footers(self, doc) -> bool:
        try:
            for section in doc.sections:
                if section.header.paragraphs or section.footer.paragraphs:
                    for p in section.header.paragraphs + section.footer.paragraphs:
                        if p.text.strip():
                            return True
        except Exception:
            pass
        return False

    # ─── Text Cleaning ───────────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace, remove control chars, collapse blank lines."""
        # Remove non-printable characters except newlines/tabs
        text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
        # Collapse multiple spaces
        text = re.sub(r" {2,}", " ", text)
        # Collapse 3+ newlines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
